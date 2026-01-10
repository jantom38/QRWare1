# -*- coding: utf-8 -*-
"""
Skrypt generujący scenariusze użycia systemu QRWare do dokumentu Word.
Wymaga instalacji: pip install python-docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


class UseCase:
    """Klasa reprezentująca scenariusz użycia."""
    
    def __init__(self, code, name, actor, goal, preconditions, postconditions,
                 main_scenario, alternative_scenarios=None):
        self.code = code
        self.name = name
        self.actor = actor
        self.goal = goal
        self.preconditions = preconditions
        self.postconditions = postconditions
        self.main_scenario = main_scenario
        self.alternative_scenarios = alternative_scenarios or []


def set_cell_shading(cell, color):
    """Ustawia kolor tła komórki."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_use_case_table(doc, use_case):
    """Dodaje tabelę scenariusza użycia do dokumentu."""
    
    # Dodaj nagłówek scenariusza
    heading = doc.add_heading(f'{use_case.code}: {use_case.name}', level=3)
    
    # Utwórz tabelę
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Ustaw szerokości kolumn
    for row in table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(12)
    
    # Wypełnij tabelę
    rows_data = [
        ('Kod', use_case.code),
        ('Aktor główny', use_case.actor),
        ('Cel', use_case.goal),
        ('Warunki wstępne', use_case.preconditions),
        ('Warunki końcowe', use_case.postconditions),
        ('Scenariusz główny', '\n'.join(f'{i+1}. {step}' for i, step in enumerate(use_case.main_scenario)))
    ]
    
    for i, (label, value) in enumerate(rows_data):
        cell_label = table.rows[i].cells[0]
        cell_value = table.rows[i].cells[1]
        
        cell_label.text = label
        cell_value.text = value
        
        # Pogrubienie etykiet
        for paragraph in cell_label.paragraphs:
            for run in paragraph.runs:
                run.bold = True
        
        # Kolor tła dla etykiet
        set_cell_shading(cell_label, "E6E6E6")
    
    # Dodaj scenariusze alternatywne jeśli istnieją
    if use_case.alternative_scenarios:
        doc.add_paragraph()
        alt_para = doc.add_paragraph()
        alt_run = alt_para.add_run('Scenariusze alternatywne:')
        alt_run.bold = True
        
        for alt in use_case.alternative_scenarios:
            doc.add_paragraph(alt, style='List Bullet')
    
    doc.add_paragraph()  # Odstęp po tabeli


def get_administrator_use_cases():
    """Zwraca scenariusze uzycia dla Administratora."""
    return [
        UseCase(
            code="UC-A01",
            name="Zarzadzanie produktami (CRUD)",
            actor="Administrator",
            goal="Utworzenie, odczyt, modyfikacja lub usuniecie produktu w systemie",
            preconditions="Administrator jest zalogowany i posiada uprawnienia do zarzadzania produktami",
            postconditions="Produkt zostal dodany/zmodyfikowany/usuniety z bazy danych",
            main_scenario=[
                "Administrator wybiera opcje 'Dodaj nowy produkt'",
                "System wyswietla formularz z polami: nazwa, SKU, kod kreskowy, kategoria, jednostka miary, cena, koszt, waga, wymiary",
                "Administrator wypelnia wymagane pola (nazwa, SKU, kategoria, jednostka miary)",
                "Administrator opcjonalnie okresla: minimalny/maksymalny stan, punkt ponownego zamowienia, warunki przechowywania, wilgotnosc",
                "Administrator potwierdza wprowadzone dane",
                "System waliduje unikalnosc SKU i poprawnosc danych",
                "System zapisuje produkt i wyswietla potwierdzenie"
            ],
            alternative_scenarios=[
                "6a. SKU juz istnieje - system wyswietla komunikat bledu i prosi o zmiane",
                "6b. Brak wymaganych pol - system podswietla brakujace pola"
            ]
        ),
        UseCase(
            code="UC-A02",
            name="Zarzadzanie kategoriami",
            actor="Administrator",
            goal="Utworzenie hierarchicznej struktury kategorii produktow",
            preconditions="Administrator jest zalogowany",
            postconditions="Kategoria zostala utworzona z odpowiednia hierarchia",
            main_scenario=[
                "Administrator wybiera opcje 'Zarzadzaj kategoriami'",
                "System wyswietla drzewiasta strukture istniejacych kategorii",
                "Administrator wybiera opcje 'Dodaj kategorie'",
                "System wyswietla formularz z polami: nazwa, kod, opis, kategoria nadrzedna, kolejnosc sortowania",
                "Administrator opcjonalnie definiuje wymagania przechowywania (temperatura min/max, wilgotnosc min/max)",
                "Administrator zapisuje kategorie",
                "System waliduje unikalnosc kodu i zapisuje kategorie w hierarchii"
            ],
            alternative_scenarios=[
                "7a. Proba usuniecia kategorii z podkategoriami lub produktami - system blokuje operacje"
            ]
        ),
        UseCase(
            code="UC-A03",
            name="Definiowanie stref magazynowych",
            actor="Administrator",
            goal="Utworzenie strefy magazynowej (np. strefa przyjecia, chlodnia, strefa wydawania)",
            preconditions="Administrator jest zalogowany",
            postconditions="Strefa magazynowa zostala utworzona w systemie",
            main_scenario=[
                "Administrator wybiera opcje 'Zarzadzaj strefami'",
                "System wyswietla liste istniejacych stref",
                "Administrator wybiera 'Dodaj strefe'",
                "System wyswietla formularz z polami: nazwa, kod, typ strefy (RECEIVING, STORAGE, COLD_STORAGE, SHIPPING, STAGING, QUARANTINE)",
                "Administrator definiuje warunki srodowiskowe (temperatura, wilgotnosc) jesli wymagane",
                "Administrator okresla pojemnosc strefy",
                "System zapisuje strefe i generuje unikalny identyfikator"
            ]
        ),
        UseCase(
            code="UC-A04",
            name="Definiowanie lokalizacji magazynowych",
            actor="Administrator",
            goal="Utworzenie lokalizacji (miejsce paletowe, regal, polka) w ramach strefy",
            preconditions="Istnieje co najmniej jedna strefa magazynowa",
            postconditions="Lokalizacja zostala przypisana do strefy",
            main_scenario=[
                "Administrator wybiera strefe docelowa",
                "Administrator wybiera opcje 'Dodaj lokalizacje'",
                "System wyswietla formularz z polami: kod lokalizacji, typ (RACK, SHELF, BIN, PALLET, FLOOR, STAGING_AREA), pojemnosc",
                "Administrator definiuje wspolrzedne (alejka, regal, poziom, pozycja)",
                "Administrator okresla ograniczenia (max waga, typy dozwolonych produktow)",
                "System waliduje unikalnosc kodu lokalizacji w ramach strefy",
                "System zapisuje lokalizacje"
            ]
        ),
        UseCase(
            code="UC-A05",
            name="Podglad historii operacji magazynowych",
            actor="Administrator",
            goal="Przegladanie historii wszystkich operacji magazynowych",
            preconditions="Administrator jest zalogowany",
            postconditions="Administrator uzyskal wglad w historie operacji",
            main_scenario=[
                "Administrator wybiera opcje 'Historia operacji'",
                "System wyswietla liste operacji z mozliwoscia filtrowania",
                "Administrator moze filtrowac po: typie operacji (RECEIPT, ISSUE, TRANSFER, ADJUSTMENT), dacie, uzytkowniku, produkcie, lokalizacji",
                "System wyswietla szczegoly: data, typ ruchu, ilosc przed/po, lokalizacja zrodlowa/docelowa, uzytkownik wykonujacy",
                "Administrator moze eksportowac raport do pliku"
            ]
        ),
        UseCase(
            code="UC-A06",
            name="Zarzadzanie rolami systemowymi",
            actor="Administrator",
            goal="Utworzenie i konfiguracja roli z zestawem uprawnien",
            preconditions="Administrator posiada uprawnienia do zarzadzania rolami",
            postconditions="Rola zostala utworzona z przypisanymi uprawnieniami",
            main_scenario=[
                "Administrator wybiera 'Zarzadzaj rolami'",
                "System wyswietla liste istniejacych rol (Administrator, Menadzer, Magazynier)",
                "Administrator wybiera 'Dodaj role'",
                "Administrator wprowadza nazwe i opis roli",
                "System wyswietla dostepne uprawnienia pogrupowane wedlug zasobow (products:create, inventory:read, orders:update, itp.)",
                "Administrator przypisuje wybrane uprawnienia do roli",
                "System zapisuje role"
            ]
        ),
        UseCase(
            code="UC-A07",
            name="Zarzadzanie uzytkownikami",
            actor="Administrator",
            goal="Utworzenie konta uzytkownika i przypisanie rol",
            preconditions="Administrator jest zalogowany z odpowiednimi uprawnieniami",
            postconditions="Uzytkownik zostal utworzony i moze sie zalogowac",
            main_scenario=[
                "Administrator wybiera 'Zarzadzaj uzytkownikami'",
                "System wyswietla liste uzytkownikow",
                "Administrator wybiera 'Dodaj uzytkownika'",
                "System wyswietla formularz: login, email, imie, nazwisko, haslo tymczasowe",
                "Administrator przypisuje role do uzytkownika",
                "System zapisuje uzytkownika i wysyla powiadomienie o utworzeniu konta",
                "Administrator moze pozniej zresetowac haslo uzytkownika"
            ]
        )
    ]


def get_manager_use_cases():
    """Zwraca scenariusze uzycia dla Menadzera."""
    return [
        UseCase(
            code="UC-M01",
            name="Tworzenie raportow o stanach magazynowych",
            actor="Menadzer",
            goal="Wygenerowanie raportu o aktualnych stanach i przeplywie produktow",
            preconditions="Menadzer jest zalogowany",
            postconditions="Raport zostal wygenerowany i wyswietlony/wyeksportowany",
            main_scenario=[
                "Menadzer wybiera 'Raporty'",
                "System wyswietla opcje raportow: stany magazynowe, niski stan, przeplywy, wartosc zapasow",
                "Menadzer wybiera typ raportu i okresla parametry (zakres dat, kategorie, strefy)",
                "System generuje raport zawierajacy: produkty z niskim stanem (ponizej reorderPoint), wartosc calkowita zapasow, produkty przeterminowane",
                "Menadzer moze wyeksportowac raport do formatu PDF/Excel"
            ]
        ),
        UseCase(
            code="UC-M02",
            name="Zarzadzanie stanami magazynowymi",
            actor="Menadzer",
            goal="Utworzenie nowego stanu magazynowego dla produktu",
            preconditions="Produkt i lokalizacja istnieja w systemie",
            postconditions="Stan magazynowy zostal utworzony z wygenerowanym kodem QR",
            main_scenario=[
                "Menadzer wybiera 'Zarzadzaj stanami magazynowymi'",
                "System wyswietla liste istniejacych stanow",
                "Menadzer wybiera 'Dodaj stan magazynowy'",
                "System wyswietla formularz: produkt, lokalizacja, ilosc, numer partii, numer serii, data przyjecia, data waznosci",
                "Menadzer moze okreslic warunki przechowywania (temperatura, wilgotnosc)",
                "System automatycznie generuje unikalny kod QR dla stanu",
                "System zapisuje stan ze statusem AVAILABLE",
                "System wyswietla potwierdzenie z kodem QR gotowym do wydruku"
            ]
        ),
        UseCase(
            code="UC-M03",
            name="Tworzenie zlecen magazynowych",
            actor="Menadzer",
            goal="Utworzenie zlecenia do realizacji przez magazyniera",
            preconditions="Menadzer jest zalogowany",
            postconditions="Zlecenie zostalo utworzone i przypisane do realizacji",
            main_scenario=[
                "Menadzer wybiera 'Utworz zlecenie'",
                "System wyswietla formularz z typami zlecen: INBOUND (przyjecie), OUTBOUND (wydanie), TRANSFER (przeniesienie), PICK (kompletacja)",
                "Menadzer wybiera typ zlecenia i okresla priorytet (LOW, NORMAL, HIGH, URGENT)",
                "Menadzer dodaje pozycje zlecenia (produkt, ilosc, lokalizacja zrodlowa/docelowa)",
                "Menadzer okresla termin realizacji (expectedDate)",
                "Menadzer opcjonalnie przypisuje zlecenie do konkretnego magazyniera",
                "System generuje unikalny numer zlecenia",
                "System zapisuje zlecenie ze statusem CREATED lub ASSIGNED"
            ]
        ),
        UseCase(
            code="UC-M04",
            name="Monitorowanie statusu zlecen",
            actor="Menadzer",
            goal="Sledzenie postepu realizacji zlecen",
            preconditions="Istnieja zlecenia w systemie",
            postconditions="Menadzer uzyskal informacje o statusie zlecen",
            main_scenario=[
                "Menadzer wybiera 'Podglad zlecen'",
                "System wyswietla dashboard ze zleceniami pogrupowanymi wedlug statusu",
                "Statusy: CREATED, ASSIGNED, IN_PROGRESS, PARTIALLY_COMPLETED, COMPLETED, CANCELLED",
                "Menadzer moze filtrowac zlecenia po: typie, priorytecie, magazynierze, dacie",
                "System wyswietla procent realizacji kazdego zlecenia (completedItems/totalItems)",
                "System oznacza zlecenia przeterminowane (isOverdue)",
                "Menadzer moze kliknac zlecenie, aby zobaczyc szczegoly i historie zmian statusu"
            ]
        ),
        UseCase(
            code="UC-M05",
            name="Generowanie kodow QR",
            actor="Menadzer",
            goal="Wygenerowanie kodu QR dla produktu lub stanu magazynowego",
            preconditions="Produkt lub stan magazynowy istnieje w systemie",
            postconditions="Kod QR zostal wygenerowany i jest gotowy do wydruku",
            main_scenario=[
                "Menadzer wybiera produkt lub stan magazynowy",
                "Menadzer wybiera opcje 'Generuj kod QR'",
                "System wyswietla opcje: typ kodu (PRODUCT, INVENTORY, LOCATION), poziom korekcji bledow, rozmiar",
                "System generuje kod QR zawierajacy: identyfikator obiektu, typ, dane pomocnicze",
                "System wyswietla podglad kodu QR",
                "Menadzer moze pobrac kod QR jako obraz PNG do wydruku",
                "Menadzer drukuje i nakleja kod na polke/palete/karton"
            ]
        ),
        UseCase(
            code="UC-M06",
            name="Skanowanie kodow QR (podglad informacji)",
            actor="Menadzer",
            goal="Uzyskanie szczegolowych informacji o produkcie/stanie poprzez skan QR",
            preconditions="Menadzer ma dostep do aplikacji mobilnej z kamera",
            postconditions="System wyswietlil szczegolowe informacje",
            main_scenario=[
                "Menadzer otwiera aplikacje Android i wybiera 'Skanuj QR'",
                "Aplikacja aktywuje kamere",
                "Menadzer skanuje kod QR na produkcie/polce",
                "System dekoduje kod i pobiera dane z bazy",
                "System wyswietla: nazwe produktu, SKU, ilosc dostepna, ilosc zarezerwowana, lokalizacje, date waznosci, historie ruchow",
                "Menadzer moze bezposrednio przejsc do edycji lub utworzenia zlecenia"
            ]
        ),
        UseCase(
            code="UC-M07",
            name="Podglad procesow logistycznych i tworzenie historii",
            actor="Menadzer",
            goal="Monitorowanie i dokumentowanie procesow zachodzacych w magazynie",
            preconditions="Menadzer jest zalogowany",
            postconditions="Historia procesow jest zapisana i dostepna do analizy",
            main_scenario=[
                "Menadzer wybiera 'Historia ruchow magazynowych'",
                "System wyswietla chronologiczna liste wszystkich operacji",
                "Dla kazdej operacji widoczne sa: typ ruchu (MovementType), data, uzytkownik, produkt, ilosc, lokalizacje",
                "Menadzer moze filtrowac po typie operacji: RECEIPT, ISSUE, TRANSFER, ADJUSTMENT, PICK, PUTAWAY",
                "System automatycznie zapisuje wszystkie operacje z sygnatura czasowa i identyfikatorem uzytkownika"
            ]
        )
    ]


def get_warehouse_worker_use_cases():
    """Zwraca scenariusze uzycia dla Magazyniera."""
    return [
        UseCase(
            code="UC-W01",
            name="Skanowanie kodow QR w celu identyfikacji produktu",
            actor="Magazynier",
            goal="Szybka identyfikacja produktu i uzyskanie informacji o stanie magazynowym",
            preconditions="Magazynier jest zalogowany do aplikacji Android",
            postconditions="Magazynier uzyskal potrzebne informacje",
            main_scenario=[
                "Magazynier otwiera aplikacje na urzadzeniu mobilnym",
                "Magazynier wybiera opcje 'Skanuj'",
                "Aplikacja uruchamia modul kamery (QrCodeAnalyzer)",
                "Magazynier skanuje kod QR na produkcie/lokalizacji",
                "System wyswietla: nazwa produktu, ilosc dostepna (availableQuantity), ilosc zarezerwowana, lokalizacja, numer partii, data waznosci",
                "System pokazuje status: AVAILABLE, RESERVED, DAMAGED, EXPIRED, QUARANTINE"
            ],
            alternative_scenarios=[
                "5a. Kod QR nierozpoznany - system wyswietla komunikat 'Nieprawidlowy kod'",
                "5b. Produkt przeterminowany - system wyswietla ostrzezenie"
            ]
        ),
        UseCase(
            code="UC-W02",
            name="Przeglad przypisanych zlecen",
            actor="Magazynier",
            goal="Wyswietlenie listy zlecen do realizacji",
            preconditions="Magazynier jest zalogowany",
            postconditions="Magazynier widzi swoje zlecenia",
            main_scenario=[
                "Magazynier wybiera 'Moje zlecenia' w aplikacji",
                "System wyswietla liste zlecen przypisanych do zalogowanego uzytkownika",
                "Zlecenia sa posortowane wedlug priorytetu (URGENT na gorze) i terminu",
                "Dla kazdego zlecenia widoczne: numer, typ, priorytet, termin, procent realizacji",
                "Magazynier moze filtrowac zlecenia po statusie",
                "System oznacza zlecenia wymagajace natychmiastowej akcji (requiresAction)"
            ]
        ),
        UseCase(
            code="UC-W03",
            name="Realizacja zlecenia kompletacji (PICK)",
            actor="Magazynier",
            goal="Skompletowanie zamowienia wedlug zlecenia",
            preconditions="Zlecenie typu PICK jest przypisane do magazyniera",
            postconditions="Zlecenie zostalo zrealizowane, stany magazynowe zaktualizowane",
            main_scenario=[
                "Magazynier wybiera zlecenie do realizacji",
                "Magazynier rozpoczyna zlecenie (status zmienia sie na IN_PROGRESS)",
                "System wyswietla liste pozycji do skompletowania z lokalizacjami zrodlowymi",
                "Dla kazdej pozycji magazynier udaje sie do wskazanej lokalizacji",
                "Magazynier skanuje kod QR lokalizacji (weryfikacja)",
                "Magazynier skanuje kod QR produktu",
                "Magazynier potwierdza pobrana ilosc",
                "System aktualizuje stan magazynowy (zmniejsza quantity, reservedQuantity)",
                "System rejestruje ruch typu ORDER_PICK w historii",
                "Po skompletowaniu wszystkich pozycji magazynier oznacza zlecenie jako zakonczone",
                "Status zlecenia zmienia sie na COMPLETED"
            ],
            alternative_scenarios=[
                "4a. Produkt niedostepny w lokalizacji - magazynier moze zglosic rozbieznosc",
                "4b. Ilosc niewystarczajaca - system oznacza zlecenie jako PARTIALLY_COMPLETED"
            ]
        ),
        UseCase(
            code="UC-W04",
            name="Realizacja zlecenia przyjecia towaru (INBOUND)",
            actor="Magazynier",
            goal="Przyjecie towaru do magazynu i umieszczenie w lokalizacji",
            preconditions="Zlecenie typu INBOUND jest aktywne",
            postconditions="Towar zostal przyjety i zarejestrowany w systemie",
            main_scenario=[
                "Magazynier wybiera zlecenie przyjecia",
                "System wyswietla oczekiwane produkty i ilosci",
                "Magazynier weryfikuje fizycznie otrzymany towar",
                "Dla kazdej pozycji magazynier skanuje/wprowadza dane produktu",
                "Magazynier potwierdza ilosc rzeczywista",
                "Magazynier skanuje kod QR docelowej lokalizacji (strefa przyjecia lub magazyn)",
                "Magazynier opcjonalnie wprowadza numer partii, date waznosci",
                "System tworzy nowy stan magazynowy (InventoryItem)",
                "System generuje kod QR dla nowego stanu",
                "System rejestruje ruch typu ORDER_RECEIPT",
                "Magazynier zamyka zlecenie"
            ],
            alternative_scenarios=[
                "3a. Rozbieznosc ilosciowa - magazynier wprowadza rzeczywista ilosc i powod roznicy"
            ]
        ),
        UseCase(
            code="UC-W05",
            name="Realizacja zlecenia wydania towaru (OUTBOUND)",
            actor="Magazynier",
            goal="Wydanie towaru z magazynu",
            preconditions="Zlecenie typu OUTBOUND jest aktywne, towar jest zarezerwowany",
            postconditions="Towar zostal wydany, stany zaktualizowane",
            main_scenario=[
                "Magazynier wybiera zlecenie wydania",
                "System wyswietla liste produktow do wydania z lokalizacjami",
                "Magazynier udaje sie do wskazanych lokalizacji",
                "Dla kazdej pozycji magazynier skanuje kod QR lokalizacji zrodlowej",
                "Magazynier skanuje kod QR produktu",
                "Magazynier potwierdza wydawana ilosc",
                "Magazynier przenosi towar do strefy wydan",
                "System zmniejsza stan magazynowy",
                "System rejestruje ruch typu ORDER_ISSUE",
                "Magazynier potwierdza zakonczenie wydania",
                "Status zlecenia zmienia sie na COMPLETED"
            ]
        ),
        UseCase(
            code="UC-W06",
            name="Realizacja przeniesienia towaru (TRANSFER)",
            actor="Magazynier",
            goal="Przeniesienie towaru miedzy lokalizacjami",
            preconditions="Zlecenie typu TRANSFER jest aktywne",
            postconditions="Towar zostal przeniesiony, lokalizacja zaktualizowana",
            main_scenario=[
                "Magazynier wybiera zlecenie przeniesienia",
                "System wyswietla: produkt, ilosc, lokalizacje zrodlowa i docelowa",
                "Magazynier skanuje kod QR w lokalizacji zrodlowej",
                "Magazynier pobiera towar",
                "Magazynier przenosi towar do lokalizacji docelowej",
                "Magazynier skanuje kod QR lokalizacji docelowej",
                "System aktualizuje lokalizacje stanu magazynowego",
                "System rejestruje ruch typu TRANSFER z fromLocation i toLocation",
                "Magazynier potwierdza zakonczenie"
            ]
        ),
        UseCase(
            code="UC-W07",
            name="Nieprzewidziane przyjecie/wydanie towaru",
            actor="Magazynier",
            goal="Zarejestrowanie nieplanowanego ruchu magazynowego",
            preconditions="Magazynier jest zalogowany",
            postconditions="Ruch zostal zarejestrowany z uzasadnieniem",
            main_scenario=[
                "Magazynier wybiera opcje 'Nieplanowane przyjecie' lub 'Nieplanowane wydanie'",
                "System wyswietla formularz",
                "Magazynier skanuje lub wybiera produkt",
                "Magazynier wprowadza ilosc i lokalizacje",
                "System wymaga podania powodu (np. zwrot od klienta, uszkodzenie, utrata, nadwyzka dostawy)",
                "Magazynier wprowadza szczegolowy opis sytuacji",
                "System tworzy lub aktualizuje stan magazynowy",
                "System rejestruje ruch typu RETURN, FOUND, DAMAGE, LOSS, SAMPLE lub DISPOSAL z uzasadnieniem",
                "System moze wymagac zatwierdzenia przez menadzera (requiresApproval)"
            ],
            alternative_scenarios=[
                "5a. Brak podania powodu - system blokuje operacje i wymaga uzupelnienia",
                "8a. Operacja wymaga zatwierdzenia - system wysyla powiadomienie do menadzera"
            ]
        )
    ]


def generate_document():
    """Generuje dokument Word ze scenariuszami uzycia."""
    
    # Utworz nowy dokument
    doc = Document()
    
    # Ustaw tytul dokumentu
    title = doc.add_heading('Scenariusze Uzycia Systemu QRWare', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Dodaj informacje wstepne
    doc.add_paragraph(
        'Niniejszy dokument zawiera szczegolowe scenariusze uzycia (Use Cases) '
        'dla systemu zarzadzania magazynem QRWare. Scenariusze sa podzielone '
        'na trzy moduly odpowiadajace rolom uzytkownikow w systemie.'
    )
    doc.add_paragraph()
    
    # Sekcja Administratora
    doc.add_heading('1. Modul Administratora', level=1)
    doc.add_paragraph(
        'Modul Administratora obejmuje funkcje zarzadzania podstawowymi obiektami '
        'systemu oraz konfiguracji uprawnien i uzytkownikow.'
    )
    for use_case in get_administrator_use_cases():
        add_use_case_table(doc, use_case)
    
    # Sekcja Menadzera
    doc.add_heading('2. Modul Menadzera', level=1)
    doc.add_paragraph(
        'Modul Menadzera obejmuje funkcje raportowania, tworzenia zlecen, '
        'generowania kodow QR oraz monitorowania procesow logistycznych.'
    )
    for use_case in get_manager_use_cases():
        add_use_case_table(doc, use_case)
    
    # Sekcja Magazyniera
    doc.add_heading('3. Modul Magazyniera', level=1)
    doc.add_paragraph(
        'Modul Magazyniera obejmuje funkcje operacyjne realizowane za pomoca '
        'aplikacji mobilnej Android, w tym skanowanie kodow QR i realizacje zlecen.'
    )
    for use_case in get_warehouse_worker_use_cases():
        add_use_case_table(doc, use_case)
    
    # Zapisz dokument
    output_path = 'Scenariusze_Uzycia_QRWare.docx'
    doc.save(output_path)
    print(f'Dokument zostal zapisany jako: {output_path}')
    return output_path


if __name__ == '__main__':
    generate_document()
